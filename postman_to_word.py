import json

from docx import Document
from docx.shared import Length, Pt, RGBColor, Cm
import demjson

RGB_COLOR_MAP = {
    "GET": RGBColor(206, 238, 197),
    "POST": RGBColor(243, 168, 83),
    "PUT": RGBColor(7, 123, 237),
    "PATCH": RGBColor(0, 0, 0),
    "DELETE": RGBColor(230, 114, 105),
}


class PostmanToWord:
    def __init__(self, file_name):
        self.document = Document()
        self.file_name = file_name

    def start_convert(self):
        with open(self.file_name) as f:
            data = json.load(f)
        level = 0
        title = data.get("info").get("name")
        items = data.get("item")
        self.document.add_heading(title, level)

        for item in items:
            self.traversal(item, level + 1)

    def traversal(self, item, level):
        # 如果是文件夹
        if item.get("item"):
            self.document.add_heading(item.get("name"), level + 1)
            for it in item.get("item"):
                self.traversal(it, level + 1)
        else:
            self.handle_request(item, level + 1)

    def handle_request(self, req, level):
        heading = req.get("name")
        request = req.get("request")
        method = request.get("method")
        url = request.get("url").get("raw")
        params = request.get("url").get("query")
        body = request.get("body", {}).get("raw")

        self.document.add_heading(heading, level)
        p = self.document.add_paragraph()
        r = p.add_run(method)
        r.font.bold = True
        r.font.color.rgb = RGB_COLOR_MAP.get(method, RGBColor(0, 0, 0))
        r = p.add_run(" " + url)
        if params:
            self.gen_param_table(params)
        if body:
            p = self.document.add_paragraph(
                "```json\n" + demjson.decode(json.dumps(body, sort_keys=True, indent=2)) + "\n```"
            )

    def gen_param_table(self, params):
        table = self.document.add_table(rows=len(params) + 1, cols=4, style="Normal Table")
        # 添加表头
        cell = table.cell(0, 0)
        p = cell.paragraphs[0]
        run = p.add_run("参数名")
        run.font.bold = True

        cell = table.cell(0, 1)
        p = cell.paragraphs[0]
        run = p.add_run("类型")
        run.font.bold = True

        cell = table.cell(0, 2)
        p = cell.paragraphs[0]
        run = p.add_run("是否必选")
        run.font.bold = True

        cell = table.cell(0, 3)
        p = cell.paragraphs[0]
        run = p.add_run("说明")
        run.font.bold = True
        # 添加参数列表
        for i in range(1, len(params) + 1):
            cell = table.cell(i, 0)
            p = cell.paragraphs[0]
            run = p.add_run(params[i - 1].get("key"))
            cell = table.cell(i, 1)
            p = cell.paragraphs[0]
            cell = table.cell(i, 2)
            p = cell.paragraphs[0]
            cell = table.cell(i, 3)
            p = cell.paragraphs[0]

    def save(self):
        self.document.save('demo.docx')


# 需要使用postman导出的collection V2.1的JSON文件
convert = PostmanToWord("postman.json")
convert.start_convert()
convert.save()
